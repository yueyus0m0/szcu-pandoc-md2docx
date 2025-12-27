from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt 
import os
import sys
import copy  # 引入深拷贝模块
import re  # 用于参考文献编号匹配

# ==========================================
#              用户配置区域 (CONFIG)
# ==========================================
CONFIG = {
    "target_file": "main.docx",           
    "output_file": "main.docx", 
    
    # --- 页眉设置 ---
    # 样式完全依赖 Reference.docx (最后一节) 的样式。
    
    # --- 页脚设置 ---
    "footer_font_name": "Times New Roman",
    "footer_font_size_pt": 9,
    "number_format": "upperRoman",       
    "start_at": "1",                     
}
# ==========================================

def sync_section_layout(source_sec, target_sec):
    """
    [物理容器同步] - 修复版
    删除不存在的 'odd_and_even_pages_header_footer' 属性。
    只同步 'different_first_page_header_footer' (首页不同)，这是节属性。
    """
    # 1. 克隆页边距 
    target_sec.top_margin = source_sec.top_margin
    target_sec.bottom_margin = source_sec.bottom_margin
    target_sec.left_margin = source_sec.left_margin
    target_sec.right_margin = source_sec.right_margin
    target_sec.gutter = source_sec.gutter
    
    # 2. 克隆页眉页脚距离 
    target_sec.header_distance = source_sec.header_distance
    target_sec.footer_distance = source_sec.footer_distance
    
    # 3. 克隆纸张大小和方向 
    target_sec.page_width = source_sec.page_width
    target_sec.page_height = source_sec.page_height
    target_sec.orientation = source_sec.orientation

    # 4. [关键修复] 同步页眉显示逻辑
    # "首页不同"是节属性，必须同步。
    # "奇偶页不同"是文档全局属性，不需要在这里同步，删除报错的那一行。
    target_sec.different_first_page_header_footer = source_sec.different_first_page_header_footer
def clone_header_xml(source_header, target_header):
    """
    [核心手术: XML 内容深拷贝]
    将源页眉的底层 XML 节点完全清空，并注入源页眉的所有子节点。
    """
    source_element = source_header._element
    target_element = target_header._element

    # 1. 清空目标页眉的所有现有内容 (使用 list() 确保安全遍历删除)
    for child in list(target_element):
        target_element.remove(child)

    # 2. 遍历源页眉的所有子节点 (Paragraphs, Tables, etc.)
    for child in source_element:
        # 使用 deepcopy 确保完全独立的副本，互不干扰
        new_child = copy.deepcopy(child)
        target_element.append(new_child)

def set_run_font(run, font_name_ascii, font_name_eastasia, size_pt):
    """(保留用于页脚) 同时设置中西文字体和字号"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name_ascii
    
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name_ascii)
    rFonts.set(qn('w:hAnsi'), font_name_ascii)
    rFonts.set(qn('w:eastAsia'), font_name_eastasia)
    rPr.append(rFonts)

def create_page_number_field(run):
    """(保留用于页脚) 插入动态页码域 {PAGE}"""
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def set_footer_and_format(section, fmt, start=None):
    """(保留用于页脚) 设置页码格式"""
    footer = section.footer
    # 清空页脚
    for p in footer.paragraphs: p.text = ""
    if not footer.paragraphs: footer.add_paragraph()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run()
    create_page_number_field(run)
    set_run_font(run, CONFIG["footer_font_name"], "宋体", CONFIG["footer_font_size_pt"])

    # 设置页码格式 (pgNumType)
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:fmt'), fmt)
    if start:
        pgNumType.set(qn('w:start'), str(start))

def apply_bibliography_styles(doc):
    """
    [参考文献样式应用]
    根据参考文献编号自动应用不同的样式：
    - 编号 10-99: 应用"书目2"样式
    - 编号 100及以后: 应用"书目3"样式
    """
    print("📚 应用参考文献样式...")
    modified_count = 0
    
    # 遍历所有段落
    for para in doc.paragraphs:
        # 检查段落是否使用"书目"样式（英文名：Bibliography）
        if para.style.name in ["书目", "Bibliography"] or "Bibliography" in para.style.name:
            # 提取段落开头的编号（支持 [1]、1、等格式）
            text = para.text.strip()
            
            # 匹配常见的参考文献编号格式
            # 例如: [1], 1., (1), 1、等
            match = re.match(r'^[\[\(]?(\d+)[\]\)\.\、]?\s', text)
            
            if match:
                number = int(match.group(1))
                new_style = None
                
                # 根据编号范围确定样式
                if 10 <= number <= 99:
                    new_style = "Bibliography 2"
                elif number >= 100:
                    new_style = "Bibliography 3"
                
                # 应用新样式
                if new_style:
                    try:
                        old_style = para.style.name
                        para.style = new_style
                        # 验证样式是否真的被应用
                        actual_style = para.style.name
                        modified_count += 1
                        print(f"   └── 编号 [{number}] -> {old_style} => {actual_style}")
                        if actual_style != new_style and actual_style != old_style:
                            print(f"       ⚠️ 警告: 样式未按预期应用！")
                    except KeyError:
                        print(f"   └── ⚠️ 警告: 样式 '{new_style}' 不存在，跳过编号 [{number}]")
    
    if modified_count > 0:
        print(f"   ✅ 成功处理 {modified_count} 个参考文献条目")
    else:
        print(f"   ⚠️ 未找到需要修改的参考文献条目")
    
    return modified_count


def run_task():
    input_path = CONFIG["target_file"]
    if not os.path.exists(input_path):
        print(f"❌ 找不到文件: {input_path}"); return

    doc = Document(input_path)
    total = len(doc.sections)
    
    # ==========================================
    # 阶段 1: 处理第一节 (受损最严重的节)
    # ==========================================
    if total > 1:
        print(f"🔧 检测到 {total} 个节，开始处理...")
        
        reference_section = doc.sections[-1] # 供体：最后一节
        target_section = doc.sections[0]     # 受体：第一节

        print("   └── [Layout] 同步第一节的物理属性(含首页显示设置)...")
        sync_section_layout(reference_section, target_section)

        print("   └── [Header] 修复第一节丢失的页眉...")
        clone_header_xml(reference_section.header, target_section.header)
    else:
        print("⚠️ 文档只有一个节，跳过布局同步。")

    print("👉 重建页脚体系 (第 1 节)...")
    sec0 = doc.sections[0]
    set_footer_and_format(sec0, fmt=CONFIG["number_format"], start=CONFIG["start_at"])

    # ==========================================
    # 阶段 2: 处理中间的节 (如果存在)
    # ==========================================
    # [关键逻辑修复]
    # 循环范围：从第 2 节开始，到倒数第 2 节结束。
    # 绝对不要包含最后一节 (total-1)，因为那是所有的样式的来源！
    
    if total > 2:
        for i in range(1, total - 1):
            sec = doc.sections[i]
            
            # 1. 强制同步版式属性 (防止中间章节因为 Pandoc 设置了首页不同而导致页眉消失)
            sec.different_first_page_header_footer = doc.sections[-1].different_first_page_header_footer
            sec.odd_and_even_pages_header_footer = doc.sections[-1].odd_and_even_pages_header_footer

            # 2. 如果页眉断开了链接，则给予修复
            if not sec.header.is_linked_to_previous:
                print(f"   └── [Header] 修复第 {i+1} 节断开的页眉...")
                clone_header_xml(doc.sections[-1].header, sec.header)
            
            # 3. 设置页脚 (阿拉伯数字)
            set_footer_and_format(sec, fmt="decimal", start="1")

    # ==========================================
    # 阶段 3: 处理最后一节 (Reference 自身)
    # ==========================================
    # 最后一节只需要重置页脚，绝对不能动它的页眉 (因为它是供体)
    if total > 1:
        print(f"👉 重建页脚体系 (最后一节)...")
        last_sec = doc.sections[-1]
        
        # [可选] 如果你希望最后一节页码顺延，把 start="1" 改为 start=None
        # 这里保持和你之前逻辑一致，强制重置为 1
        set_footer_and_format(last_sec, fmt="decimal", start="1") 

    # ==========================================
    # 阶段 4: 应用参考文献样式
    # ==========================================
    print("👉 应用参考文献样式...")
    apply_bibliography_styles(doc)

    doc.save(CONFIG["output_file"])
    print(f"✅ 自动化排版完成！文件已保存至: {CONFIG['output_file']}")

if __name__ == '__main__':
    # CLI 参数支持
    # 用法: python fix_word_layout.py [input_file] [output_file]
    if len(sys.argv) > 1:
        CONFIG["target_file"] = sys.argv[1]
        
        if len(sys.argv) > 2:
            CONFIG["output_file"] = sys.argv[2]
        else:
            # 如果只提供输入文件，默认覆盖输出
            CONFIG["output_file"] = sys.argv[1]
            
    print(f"📄 Input:  {CONFIG['target_file']}")
    print(f"📤 Output: {CONFIG['output_file']}")
    
    run_task()